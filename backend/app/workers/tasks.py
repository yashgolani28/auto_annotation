import os
import json
import math
import re
import time
import random
import shutil
import hashlib
import platform
from pathlib import Path
from datetime import datetime
from typing import Any

from sqlalchemy.orm import Session

import torch
from ultralytics import YOLO

from app.workers.celery_app import celery
from app.db.session import SessionLocal
from app.models.models import Job, Dataset, DatasetItem, ModelWeight, LabelClass, AnnotationSet, Annotation
from app.core.config import settings
from app.services.inference import load_ultralytics_model, predict_bboxes


# ---------------------------------------------------------------------
# utils
# ---------------------------------------------------------------------

def _update_job(
    db: Session,
    job: Job,
    status: str | None = None,
    progress: float | None = None,
    message: str | None = None,
):
    if status is not None:
        job.status = status
    if progress is not None:
        job.progress = float(progress)
    if message is not None:
        job.message = message
    job.updated_at = datetime.utcnow()
    db.add(job)
    db.commit()


def _sha16(s: str) -> str:
    return hashlib.sha256(s.encode("utf-8")).hexdigest()[:16]


def _safe_posix(p: Path) -> str:
    return p.as_posix().replace("\\", "/")


def _ensure_empty_dir(p: Path):
    if p.exists():
        shutil.rmtree(p, ignore_errors=True)
    p.mkdir(parents=True, exist_ok=True)


def _hardlink_or_copy(src: Path, dst: Path) -> None:
    dst.parent.mkdir(parents=True, exist_ok=True)
    if dst.exists():
        return
    try:
        os.link(src, dst)
    except Exception:
        shutil.copy2(src, dst)


def _write_yolo_data_yaml(dst: Path, dataset_root: Path, names: list[str]) -> None:
    lines = []
    lines.append(f"train: {_safe_posix(dataset_root / 'images' / 'train')}")
    lines.append(f"val:   {_safe_posix(dataset_root / 'images' / 'val')}")
    lines.append(f"test:  {_safe_posix(dataset_root / 'images' / 'test')}")
    lines.append("")
    lines.append(f"nc: {len(names)}")
    lines.append("names:")
    for nm in names:
        lines.append(f"  - '{nm}'")
    dst.write_text("\n".join(lines) + "\n", encoding="utf-8")


def _random_split_items(items: list, seed: int, train_ratio: float, val_ratio: float, test_ratio: float):
    assert abs((train_ratio + val_ratio + test_ratio) - 1.0) < 1e-6
    rng = random.Random(seed)
    idxs = list(range(len(items)))
    rng.shuffle(idxs)
    n = len(items)
    n_train = int(math.floor(train_ratio * n))
    n_val = int(math.floor(val_ratio * n))
    train_ids = set(idxs[:n_train])
    val_ids = set(idxs[n_train:n_train + n_val])
    for i, it in enumerate(items):
        if i in train_ids:
            it.split = "train"
        elif i in val_ids:
            it.split = "val"
        else:
            it.split = "test"


def _save_text(p: Path, s: str):
    p.parent.mkdir(parents=True, exist_ok=True)
    p.write_text(s, encoding="utf-8")


def _cleanup_unwanted_artifacts(artifacts_dir: Path) -> None:
    """
    If a previous run created exports (onnx/engine/tf/paddle/ncnn/...) remove them,
    because you only want model.pt + bench_runs + report.
    """
    if not artifacts_dir.exists():
        return

    kill = [
        "model.onnx",
        "model.engine",
        "model.torchscript",
        "model.mlpackage",
        "model_openvino_model",
        "model_saved_model",
        "model_paddle_model",
        "model_ncnn_model",
        "model_mnn_model",
        "model_tflite",
        "model_edgetpu.tflite",
        "model_web_model",
        "model_coreml_model",
    ]
    for name in kill:
        p = artifacts_dir / name
        try:
            if p.is_dir():
                shutil.rmtree(p, ignore_errors=True)
            elif p.exists():
                p.unlink(missing_ok=True)
        except Exception:
            pass

    # also delete obvious ultralytics export leftovers
    for p in artifacts_dir.glob("*.onnx"):
        try:
            p.unlink(missing_ok=True)
        except Exception:
            pass
    for p in artifacts_dir.glob("*.engine"):
        try:
            p.unlink(missing_ok=True)
        except Exception:
            pass


# ---------------------------------------------------------------------
# DB -> YOLO training dataset export (kept from your robust version)
# ---------------------------------------------------------------------

def _export_training_yolo_from_db(
    *,
    db: Session,
    project_id: int,
    dataset_id: int,
    annotation_set_id: int,
    approved_only: bool,
    out_root: Path,
    names_in_order: list[str],
    job: Job,
):
    _ensure_empty_dir(out_root)
    (out_root / "images" / "train").mkdir(parents=True, exist_ok=True)
    (out_root / "images" / "val").mkdir(parents=True, exist_ok=True)
    (out_root / "images" / "test").mkdir(parents=True, exist_ok=True)
    (out_root / "labels" / "train").mkdir(parents=True, exist_ok=True)
    (out_root / "labels" / "val").mkdir(parents=True, exist_ok=True)
    (out_root / "labels" / "test").mkdir(parents=True, exist_ok=True)

    items = (
        db.query(DatasetItem)
        .filter(DatasetItem.dataset_id == dataset_id)
        .order_by(DatasetItem.id.asc())
        .all()
    )
    if not items:
        raise RuntimeError("dataset has no items")

    classes = (
        db.query(LabelClass)
        .filter(LabelClass.project_id == project_id)
        .order_by(LabelClass.order_index.asc())
        .all()
    )
    if not classes:
        raise RuntimeError("no label classes in project")
    class_id_to_idx = {c.id: i for i, c in enumerate(classes)}

    q = db.query(Annotation).filter(Annotation.annotation_set_id == annotation_set_id)
    if approved_only:
        q = q.filter(Annotation.approved == True)  # noqa: E712
    anns = q.all()

    by_item: dict[int, list[Annotation]] = {}
    for a in anns:
        by_item.setdefault(a.dataset_item_id, []).append(a)

    total = len(items)
    kept = 0

    for idx, it in enumerate(items, start=1):
        img_src = Path(settings.storage_dir) / it.rel_path
        if not img_src.exists():
            continue

        split = (it.split or "train").lower()
        if split not in ("train", "val", "test"):
            split = "train"

        img_dst = out_root / "images" / split / it.file_name
        _hardlink_or_copy(img_src, img_dst)

        lbl_dst = out_root / "labels" / split / f"{Path(it.file_name).stem}.txt"

        item_anns = by_item.get(it.id, [])
        if not item_anns:
            lbl_dst.write_text("", encoding="utf-8")
        else:
            lines: list[str] = []
            for a in item_anns:
                cls_i = class_id_to_idx.get(a.class_id)
                if cls_i is None:
                    continue
                x_c = (a.x + (a.w / 2.0)) / float(it.width)
                y_c = (a.y + (a.h / 2.0)) / float(it.height)
                w_n = a.w / float(it.width)
                h_n = a.h / float(it.height)

                x_c = min(max(x_c, 0.0), 1.0)
                y_c = min(max(y_c, 0.0), 1.0)
                w_n = min(max(w_n, 0.0), 1.0)
                h_n = min(max(h_n, 0.0), 1.0)
                lines.append(f"{cls_i} {x_c:.6f} {y_c:.6f} {w_n:.6f} {h_n:.6f}")

            lbl_dst.write_text("\n".join(lines) + ("\n" if lines else ""), encoding="utf-8")

        kept += 1
        if idx % 25 == 0:
            _update_job(
                db, job,
                progress=min(0.20, (idx / total) * 0.20),
                message=f"prepared labels {idx}/{total}"
            )

    _write_yolo_data_yaml(out_root / "data.yaml", out_root, names_in_order)

    return {
        "items_total": total,
        "items_kept": kept,
        "approved_only": approved_only,
        "annotation_set_id": annotation_set_id,
        "dataset_id": dataset_id,
    }


def _count_label_lines(lbl_dir: Path) -> int:
    if not lbl_dir.exists():
        return 0
    total = 0
    for p in lbl_dir.glob("*.txt"):
        try:
            s = p.read_text(encoding="utf-8").strip()
            if s:
                total += len([ln for ln in s.splitlines() if ln.strip()])
        except Exception:
            pass
    return total


# ---------------------------------------------------------------------
# Word report generation
# ---------------------------------------------------------------------

_PLOT_CANDIDATES = [
    "results.png",
    "confusion_matrix.png",
    "confusion_matrix_normalized.png",
    "BoxPR_curve.png",
    "BoxF1_curve.png",
    "BoxP_curve.png",
    "BoxR_curve.png",
    "val_batch0_labels.jpg",
    "val_batch0_pred.jpg",
    "val_batch1_labels.jpg",
    "val_batch1_pred.jpg",
    "val_batch2_labels.jpg",
    "val_batch2_pred.jpg",
]

def _slug_name(s: str) -> str:
    s = (s or "").strip()
    if not s:
        return "trained_model"
    s = re.sub(r"\s+", "_", s)
    s = re.sub(r"[^a-zA-Z0-9._-]+", "", s)
    s = s.strip("._-")
    if not s:
        return "trained_model"
    return s[:80]

def _write_benchmark_docx_report(
    *,
    out_path: Path,
    job: Job,
    model_path: Path,
    framework: str,
    device: str,
    env: dict[str, Any],
    dataset_info: dict[str, Any],
    metrics: dict[str, Any],
    plots_dir: Path,
    notes: list[str] | None = None,
) -> tuple[bool, str]:
    """
    Returns (ok, message). If python-docx is missing, returns (False, reason).
    """
    try:
        from docx import Document  # type: ignore
        from docx.shared import Inches  # type: ignore
    except Exception as e:
        return False, f"python-docx not installed: {e}"

    doc = Document()

    doc.add_heading("yolo training benchmark report", level=1)

    # header block
    doc.add_paragraph(f"job: {job.id}")
    doc.add_paragraph(f"model: {str(model_path)}")
    doc.add_paragraph(f"date evaluated (utc): {datetime.utcnow().strftime('%Y-%m-%d %H:%M:%S')}")
    doc.add_paragraph(f"framework: {framework}")
    doc.add_paragraph(f"device: {device}")

    doc.add_heading("overview", level=2)
    doc.add_paragraph(
        "this report summarizes evaluation metrics, confusion analysis, and validation plots for the trained model."
    )

    doc.add_heading("dataset information", level=2)
    for k in ["dataset_id", "annotation_set_id", "data_yaml", "imgsz", "split"]:
        if k in dataset_info:
            doc.add_paragraph(f"{k}: {dataset_info[k]}")

    # classes
    classes = dataset_info.get("classes", [])
    if classes:
        doc.add_paragraph(f"classes ({len(classes)} total): " + ", ".join(classes))

    doc.add_heading("validation metrics", level=2)
    # simple key-value list (keeps it stable even if ultralytics keys change)
    for k, v in metrics.items():
        doc.add_paragraph(f"{k}: {v}")

    doc.add_heading("confusion matrix + curves", level=2)
    inserted_any = False
    for name in _PLOT_CANDIDATES:
        p = plots_dir / name
        if p.exists():
            inserted_any = True
            doc.add_paragraph(name)
            # fit nicely on page
            doc.add_picture(str(p), width=Inches(6.2))

    if not inserted_any:
        doc.add_paragraph("no plots were found in the validation output directory.")

    if notes:
        doc.add_heading("notes", level=2)
        for n in notes:
            doc.add_paragraph(f"- {n}")

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out_path)
    return True, "ok"


# ---------------------------------------------------------------------
# auto annotate (your working version kept; only tiny safety tweaks)
# ---------------------------------------------------------------------

@celery.task(name="auto_annotate_task")
def auto_annotate_task(job_id: int):
    db = SessionLocal()
    job = None
    try:
        job = db.query(Job).filter(Job.id == job_id).first()
        if not job:
            return

        _update_job(db, job, status="running", progress=0.0, message="starting")

        payload = job.payload or {}
        model_id = int(payload.get("model_id") or payload.get("model_weight_id", 0))
        dataset_id = int(payload["dataset_id"])
        conf = float(payload.get("conf", 0.25))
        iou = float(payload.get("iou", 0.5))
        device = str(payload.get("device", "") or "")

        # device fallback
        if device.lower() != "cpu":
            try:
                if not torch.cuda.is_available():
                    _update_job(db, job, message="cuda not available, falling back to cpu")
                    device = "cpu"
            except Exception:
                device = "cpu"

        ds = db.query(Dataset).filter(Dataset.id == dataset_id, Dataset.project_id == job.project_id).first()
        if not ds:
            _update_job(db, job, status="failed", message="dataset not found")
            return

        mw = db.query(ModelWeight).filter(ModelWeight.id == model_id, ModelWeight.project_id == job.project_id).first()
        if not mw:
            _update_job(db, job, status="failed", message="model not found")
            return

        weights_path = Path(settings.storage_dir) / mw.rel_path
        if mw.framework != "ultralytics" or not weights_path.exists():
            _update_job(db, job, status="failed", message="auto annotate supports only ultralytics .pt in this MVP")
            return

        model = load_ultralytics_model(weights_path)

        items = (
            db.query(DatasetItem)
            .filter(DatasetItem.dataset_id == dataset_id)
            .order_by(DatasetItem.id.asc())
            .all()
        )
        total = len(items)
        if total == 0:
            _update_job(db, job, status="success", progress=1.0, message="no items")
            return

        # Create annotation set
        aset = AnnotationSet(
            project_id=job.project_id,
            name=f"auto_run_{datetime.utcnow().strftime('%Y%m%d_%H%M%S')}",
            source="auto",
            model_weight_id=mw.id,
        )
        db.add(aset)
        db.commit()
        db.refresh(aset)

        classes = db.query(LabelClass).filter(LabelClass.project_id == job.project_id).all()
        if not classes:
            _update_job(db, job, status="failed", message="no label classes defined in project. please add classes first.")
            return
        name_to_class_id = {c.name.lower(): c.id for c in classes}

        # wipe any existing for this set
        db.query(Annotation).filter(Annotation.annotation_set_id == aset.id).delete(synchronize_session=False)
        db.commit()

        annotations_created = 0
        skipped_no_match = 0

        for idx, it in enumerate(items, start=1):
            img_path = Path(settings.storage_dir) / it.rel_path
            if not img_path.exists():
                continue

            try:
                preds = predict_bboxes(model, img_path, conf=conf, iou=iou, device=device)
            except Exception as e:
                _update_job(db, job, progress=idx / total, message=f"error processing {idx}/{total}: {e}")
                continue

            for p in preds:
                cls_idx = int(p["cls_idx"])
                try:
                    cls_name = str(model.names[cls_idx]).lower()
                except Exception:
                    cls_name = str(cls_idx)

                target_id = name_to_class_id.get(cls_name)
                if not target_id:
                    skipped_no_match += 1
                    continue

                x, y, w, h = p["xywh"]
                db.add(Annotation(
                    annotation_set_id=aset.id,
                    dataset_item_id=it.id,
                    class_id=target_id,
                    x=float(x), y=float(y), w=float(w), h=float(h),
                    confidence=float(p["conf"]) if p.get("conf") is not None else None,
                    approved=False,
                ))
                annotations_created += 1

            db.commit()
            _update_job(db, job, progress=idx / total, message=f"processed {idx}/{total} ({annotations_created} created)")

        msg = f"done. created {annotations_created} annotations"
        if skipped_no_match:
            msg += f", skipped {skipped_no_match} unmatched predictions"
        _update_job(db, job, status="success", progress=1.0, message=msg)

    except Exception as e:
        if job is not None:
            try:
                _update_job(db, job, status="failed", message=str(e))
            except Exception:
                pass
    finally:
        db.close()


# ---------------------------------------------------------------------
# train yolo (FIXED: no yolo benchmark cli, docx report, metadata embedding, cleanup)
# ---------------------------------------------------------------------

# --- torch.save legacy serialization (helps Windows + some environments) ---
_ORIGINAL_TORCH_SAVE = torch.save
def _torch_save_legacy(*args, **kwargs):
    if "_use_new_zipfile_serialization" not in kwargs:
        kwargs["_use_new_zipfile_serialization"] = False
    return _ORIGINAL_TORCH_SAVE(*args, **kwargs)
torch.save = _torch_save_legacy
# ------------------------------------------------------------------------


@celery.task(name="train_yolo_task")
def train_yolo_task(job_id: int):
    db = SessionLocal()
    job = None
    try:
        job = db.query(Job).filter(Job.id == job_id).first()
        if not job:
            return
        _update_job(db, job, status="running", progress=0.0, message="starting training job")

        payload = job.payload or {}

        dataset_id = int(payload["dataset_id"])
        annotation_set_id = int(payload["annotation_set_id"])
        base_model_id = int(payload["base_model_id"])
        trained_model_name = _slug_name(str(payload.get("trained_model_name", "trained_model")))
        
        split_mode = str(payload.get("split_mode", "keep"))
        train_ratio = float(payload.get("train_ratio", 0.8))
        val_ratio = float(payload.get("val_ratio", 0.1))
        test_ratio = float(payload.get("test_ratio", 0.1))
        seed = int(payload.get("seed", 1337))

        imgsz = int(payload.get("imgsz", 640))
        epochs = int(payload.get("epochs", 50))
        batch = int(payload.get("batch", 16))
        device = str(payload.get("device", "0"))

        # IMPORTANT: keep workers=0 in containers to avoid dataloader worker crashes
        # (you already observed "DataLoader worker exited unexpectedly")
        workers = 0

        optimizer = str(payload.get("optimizer", "SGD"))
        cos_lr = bool(payload.get("cos_lr", True))
        patience = int(payload.get("patience", 20))
        cache = payload.get("cache", "disk")
        approved_only = bool(payload.get("approved_only", True))

        bench_split = str(payload.get("bench_split", "test"))
        conf = float(payload.get("conf", 0.25))
        iou = float(payload.get("iou", 0.7))
        user_meta = payload.get("meta", {}) or {}

        # device fallback
        if device.lower() != "cpu":
            try:
                if not torch.cuda.is_available():
                    _update_job(db, job, message="cuda not available in this runtime, falling back to cpu")
                    device = "cpu"
            except Exception:
                device = "cpu"

        ds = db.query(Dataset).filter(Dataset.id == dataset_id, Dataset.project_id == job.project_id).first()
        if not ds:
            _update_job(db, job, status="failed", message="dataset not found")
            return

        aset = db.query(AnnotationSet).filter(
            AnnotationSet.id == annotation_set_id,
            AnnotationSet.project_id == job.project_id
        ).first()
        if not aset:
            _update_job(db, job, status="failed", message="annotation set not found")
            return

        base_mw = db.query(ModelWeight).filter(
            ModelWeight.id == base_model_id,
            ModelWeight.project_id == job.project_id
        ).first()
        if not base_mw:
            _update_job(db, job, status="failed", message="base model not found")
            return

        # prepare split if requested
        items = db.query(DatasetItem).filter(DatasetItem.dataset_id == dataset_id).order_by(DatasetItem.id.asc()).all()
        if not items:
            _update_job(db, job, status="failed", message="dataset has no items")
            return
        if split_mode == "random":
            _random_split_items(items, seed=seed, train_ratio=train_ratio, val_ratio=val_ratio, test_ratio=test_ratio)
            for it in items:
                db.add(it)
            db.commit()

        # names in order_index order
        classes = (
            db.query(LabelClass)
            .filter(LabelClass.project_id == job.project_id)
            .order_by(LabelClass.order_index.asc())
            .all()
        )
        names_in_order = [c.name for c in classes]

        # storage locations
        base_dir = Path(settings.storage_dir) / "trainings" / f"job_{job.id}"
        dataset_dir = base_dir / "dataset"
        runs_dir = base_dir / "runs"
        artifacts_dir = base_dir / "artifacts"
        artifacts_dir.mkdir(parents=True, exist_ok=True)

        # cleanup any old exports
        _cleanup_unwanted_artifacts(artifacts_dir)

        # export YOLO dataset from DB (20% progress cap)
        export_manifest = _export_training_yolo_from_db(
            db=db,
            project_id=job.project_id,
            dataset_id=dataset_id,
            annotation_set_id=annotation_set_id,
            approved_only=approved_only,
            out_root=dataset_dir,
            names_in_order=names_in_order,
            job=job,
        )
        _update_job(db, job, progress=0.20, message="dataset prepared")

        labels_train = _count_label_lines(dataset_dir / "labels" / "train")
        labels_val = _count_label_lines(dataset_dir / "labels" / "val")
        labels_test = _count_label_lines(dataset_dir / "labels" / "test")
        labels_total = labels_train + labels_val + labels_test

        if labels_total == 0:
            _update_job(
                db, job, status="failed",
                message=(
                    "exported dataset has 0 labels (all images are background). "
                    "this usually happens when approved_only=true but annotations are not approved, "
                    "or the selected annotation set has no annotations."
                )
            )
            return

        base_weights_path = Path(settings.storage_dir) / base_mw.rel_path
        if not base_weights_path.exists():
            _update_job(db, job, status="failed", message="base weights file missing in storage")
            return

        if device.lower() == "cpu":
            os.environ["CUDA_VISIBLE_DEVICES"] = "-1"

        run_name = f"train_{datetime.utcnow().strftime('%Y%m%d_%H%M%S')}_{_sha16(json.dumps(payload, sort_keys=True))}"
        _update_job(db, job, progress=0.22, message=f"training started ({run_name})")

        model = YOLO(str(base_weights_path))

        def _train_with_batch(cur_batch: int):
            return model.train(
                data=str(dataset_dir / "data.yaml"),
                imgsz=imgsz,
                epochs=epochs,
                batch=cur_batch,
                device=device,
                project=str(runs_dir),
                name=run_name,
                workers=workers,
                patience=patience,
                cos_lr=cos_lr,
                save=True,
                cache=cache,
                pretrained=True,
                optimizer=optimizer,
                verbose=True,
                seed=seed,
                amp=True,
            )

        # OOM-safe batch fallback
        cur_batch = batch
        while True:
            try:
                _train_with_batch(cur_batch)
                break
            except RuntimeError as e:
                msg = str(e).lower()
                if ("out of memory" in msg) or ("cuda" in msg and "memory" in msg):
                    if cur_batch <= 2:
                        raise
                    cur_batch = max(2, cur_batch // 2)
                    _update_job(db, job, message=f"cuda oom: retrying with batch={cur_batch}")
                    try:
                        torch.cuda.empty_cache()
                    except Exception:
                        pass
                    continue
                raise

        _update_job(db, job, progress=0.78, message="training finished, collecting artifacts")

        run_dir = runs_dir / run_name
        weights_dir = run_dir / "weights"
        best_pt = weights_dir / "best.pt"
        last_pt = weights_dir / "last.pt"
        src_ckpt = best_pt if best_pt.exists() else last_pt if last_pt.exists() else None
        if not src_ckpt:
            _update_job(db, job, status="failed", message="training completed but best.pt/last.pt not found")
            return

        # copy output model to artifacts
        out_model_path = artifacts_dir / "model.pt"
        shutil.copyfile(src_ckpt, out_model_path)

        # embed ESSI metadata into checkpoint (compatible with Ultralytics)
        ckpt = torch.load(out_model_path, map_location="cpu")
        essi_meta = {
            "job_id": job.id,
            "trained_at_utc": datetime.utcnow().isoformat(),
            "trained_model_name": trained_model_name,
            "project_id": job.project_id,
            "dataset_id": dataset_id,
            "annotation_set_id": annotation_set_id,
            "base_model_id": base_model_id,
            "export_manifest": export_manifest,
            "train": {
                "imgsz": imgsz,
                "epochs": epochs,
                "batch": cur_batch,
                "device": device,
                "workers": workers,
                "optimizer": optimizer,
                "cos_lr": cos_lr,
                "patience": patience,
                "cache": cache,
                "seed": seed,
                "split_mode": split_mode,
                "ratios": {"train": train_ratio, "val": val_ratio, "test": test_ratio},
            },
            "bench": {"split": bench_split, "conf": conf, "iou": iou},
            "user_meta": user_meta,
        }
        if isinstance(ckpt, dict):
            ckpt["essi_meta"] = essi_meta
            torch.save(ckpt, out_model_path)

        _update_job(db, job, progress=0.82, message="benchmarking: running ultralytics val (plots enabled)")

        trained = YOLO(str(out_model_path))
        metrics_obj = trained.val(
            data=str(dataset_dir / "data.yaml"),
            split=bench_split,
            imgsz=imgsz,
            device=device,
            conf=conf,
            iou=iou,
            plots=True,
            save_json=True,
            verbose=False,
            project=str(artifacts_dir / "bench_runs"),
            name=f"val_{bench_split}",
            workers=0,
        )

        # Where Ultralytics saved plots
        val_dir = Path(str(getattr(metrics_obj, "save_dir", ""))) if hasattr(metrics_obj, "save_dir") else (artifacts_dir / "bench_runs" / f"val_{bench_split}")

        results_dict = getattr(metrics_obj, "results_dict", {}) or {}
        # normalize common keys you show in UI
        prec = results_dict.get("metrics/precision(B)", results_dict.get("precision(B)", None))
        rec = results_dict.get("metrics/recall(B)", results_dict.get("recall(B)", None))
        map50 = results_dict.get("metrics/mAP50(B)", results_dict.get("mAP50(B)", None))
        map5095 = results_dict.get("metrics/mAP50-95(B)", results_dict.get("mAP50-95(B)", None))

        # Write Word report (or markdown fallback)
        _update_job(db, job, progress=0.93, message="writing benchmark report (docx)")

        ul_ver = "unknown"
        try:
            import ultralytics  # type: ignore
            ul_ver = ultralytics.__version__
        except Exception:
            pass

        env = {
            "ultralytics_version": ul_ver,
            "torch_version": getattr(torch, "__version__", "unknown"),
            "python": platform.python_version(),
            "os": f"{platform.system()} {platform.release()}",
        }

        dataset_info = {
            "dataset_id": dataset_id,
            "annotation_set_id": annotation_set_id,
            "data_yaml": _safe_posix(dataset_dir / "data.yaml"),
            "imgsz": f"{imgsz}x{imgsz}",
            "split": bench_split,
            "classes": names_in_order,
        }

        metrics_simple = {
            "precision(B)": prec,
            "recall(B)": rec,
            "mAP50(B)": map50,
            "mAP50-95(B)": map5095,
            "val_dir": _safe_posix(val_dir),
        }

        report_docx = artifacts_dir / "benchmark_report.docx"
        ok, reason = _write_benchmark_docx_report(
            out_path=report_docx,
            job=job,
            model_path=out_model_path,
            framework=f"ultralytics yolo {ul_ver}",
            device=device,
            env=env,
            dataset_info=dataset_info,
            metrics=metrics_simple,
            plots_dir=val_dir,
            notes=[
                "this report includes ultralytics validation plots (results + curves) and confusion matrix when available.",
                "extra export formats are intentionally disabled in this pipeline to avoid toolchain crashes.",
            ],
        )

        if not ok:
            # fallback markdown so job still succeeds (and tells you how to fix)
            report_md = artifacts_dir / "benchmark_report.md"
            md = []
            md.append("# yolo training benchmark report")
            md.append("")
            md.append(f"**job:** `{job.id}`  ")
            md.append(f"**model:** `{_safe_posix(out_model_path)}`  ")
            md.append(f"**date:** `{datetime.utcnow().strftime('%Y-%m-%d %H:%M:%S UTC')}`  ")
            md.append(f"**device:** `{device}`  ")
            md.append("")
            md.append("## metrics (ultralytics val)")
            md.append("")
            for k, v in metrics_simple.items():
                md.append(f"- {k}: `{v}`")
            md.append("")
            md.append("## plots dir")
            md.append(f"`{_safe_posix(val_dir)}`")
            md.append("")
            md.append("## docx report")
            md.append(f"_not generated: {reason}_")
            md.append("")
            md.append("install python-docx in backend container (requirements.txt) to enable docx.")
            _save_text(report_md, "\n".join(md) + "\n")

        # Final cleanup: ensure no exports exist
        _cleanup_unwanted_artifacts(artifacts_dir)

        # Create ModelWeight DB row
        rel_model_path = str(out_model_path.relative_to(Path(settings.storage_dir))).replace("\\", "/")
        rel_report_path = None
        if report_docx.exists():
            rel_report_path = str(report_docx.relative_to(Path(settings.storage_dir))).replace("\\", "/")
        else:
            # markdown fallback
            md_path = artifacts_dir / "benchmark_report.md"
            if md_path.exists():
                rel_report_path = str(md_path.relative_to(Path(settings.storage_dir))).replace("\\", "/")

        meta = {
            "type": "trained",
            "trained_at_utc": datetime.utcnow().isoformat(),
            "trained_model_name": trained_model_name,
            "dataset_id": dataset_id,
            "annotation_set_id": annotation_set_id,
            "base_model_id": base_model_id,
            "export_manifest": export_manifest,
            "train": {
                "imgsz": imgsz, "epochs": epochs, "batch": cur_batch, "device": device,
                "workers": workers, "optimizer": optimizer, "cos_lr": cos_lr, "patience": patience,
                "cache": cache, "seed": seed, "split_mode": split_mode,
                "ratios": {"train": train_ratio, "val": val_ratio, "test": test_ratio},
            },
            "bench": {
                "split": bench_split,
                "conf": conf,
                "iou": iou,
                "precision(B)": prec,
                "recall(B)": rec,
                "mAP50(B)": map50,
                "mAP50-95(B)": map5095,
                "val_dir_rel_path": str(val_dir.relative_to(Path(settings.storage_dir))).replace("\\", "/")
                if str(val_dir).startswith(str(settings.storage_dir)) else str(val_dir),
                "benchmark_report_rel_path": rel_report_path,
                "exports_disabled": True,
            },
            "env": env,
            "user_meta": user_meta,
        }

        mw = ModelWeight(
            project_id=job.project_id,
            name=trained_model_name,
            framework="ultralytics",
            rel_path=rel_model_path,
            class_names={str(i): n for i, n in enumerate(names_in_order)},
            meta=meta,
        )
        db.add(mw)
        db.commit()
        db.refresh(mw)

        job.payload = {
            **payload,
            "trained_model_id": mw.id,
            "benchmark_report_rel_path": rel_report_path,
            "trained_model_name": trained_model_name,
        }
        _update_job(
            db, job,
            status="success",
            progress=1.0,
            message=(
                f"done. model saved as id={mw.id}. "
                f"kept artifacts: model.pt, bench_runs, benchmark_report.{ 'docx' if report_docx.exists() else 'md' }"
            ),
        )

    except Exception as e:
        try:
            job = db.query(Job).filter(Job.id == job_id).first()
            if job:
                _update_job(db, job, status="failed", message=str(e))
        except Exception:
            pass
    finally:
        db.close()

